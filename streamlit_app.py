import streamlit as st
import pandas as pd
import openai
import os
from docx import Document
from docx.shared import Inches


# Activar el wide mode
st.set_page_config(layout="wide")

# Accedemos a la clave de API de OpenAI a través de una variable de entorno
openai.api_key = os.environ.get("OPENAI_API_KEY")

# Agregamos un título al principio
st.title('Evaluador de ensayos')

# Agregamos información de instrucciones
st.write('Suba un archivo .XLSX con los ensayos de sus alumnos. Máximo: 10 ensayos.')

# Pedimos al usuario que suba el archivo Excel
archivo = st.file_uploader('Cargar archivo Excel', type=['xlsx'])

if archivo:
    # Leemos el archivo con pandas
    data = pd.read_excel(archivo)

    # Pedimos al usuario que seleccione las columnas con el título y el ensayo
    columnas = data.columns
    columna_titulo = st.selectbox('Selecciona la columna que contiene los títulos:', columnas)
    columna_ensayo = st.selectbox('Selecciona la columna que contiene los ensayos:', columnas)

    # Agregamos un botón para iniciar la evaluación
    if st.button('Evaluar'):
        # Obtenemos los títulos y los ensayos del archivo
        titulos = data[columna_titulo].head(10).tolist()
        ensayos = data[columna_ensayo].head(10).tolist()

        # Utilizamos la API de GPT-3 para calificar cada ensayo
        resultados = []
        for i, ensayo in enumerate(ensayos):
            prompt = f"Califica el ensayo titulado '{titulos[i]}'. "
            prompt += f"Ensayo: {ensayo}. "
            response = openai.Completion.create(
                engine="text-davinci-003",
                prompt=prompt,
                temperature=0.5,
                max_tokens=512,
                n=1,
                stop=None
            )
            justificacion = response.choices[0].text.strip()

            # Agregamos sugerencias de mejora a la justificación
            response = openai.Completion.create(
                engine="text-davinci-003",
                prompt=f"Sugiere mejoras para el ensayo titulado '{titulos[i]}'. Ensayo: {ensayo}",
                temperature=0,
                max_tokens=512,
                n=1,
                stop=None,
                timeout=60,
            )
            sugerencias = response.choices[0].text.strip()

            # Agregamos la calificación y las sugerencias de mejora a la tabla
            resultados.append({
                'Ensayo': titulos[i],
                'Justificación': justificacion,
                'Sugerencias de mejora': sugerencias,
            })

        # Mostramos los resultados en una tabla
        st.write('Resultados:')
        tabla = pd.DataFrame(resultados)
        st.table(tabla)

        # Agregamos un botón para exportar los resultados a un archivo Word
        if st.button('Exportar resultados a Word'):
            st.write('Exportando resultados...')
            with st.spinner('Exportando resultados...'):
        # Creamos un nuevo documento Word
        document = Document()

        # Agregamos un encabezado
        document.add_heading('Resultados de la evaluación', 0)

        # Agregamos la tabla de resultados
        table = document.add_table(rows=len(resultados)+1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Ensayo'
        hdr_cells[1].text = 'Justificación'
        hdr_cells[2].text = 'Sugerencias de mejora'
        for i, resultado in enumerate(resultados):
            row_cells = table.rows[i+1].cells
            row_cells[0].text = resultado['Ensayo']
            row_cells[1].text = resultado['Justificación']
            row_cells[2].text = resultado['Sugerencias de mejora']

        # Guardamos el documento en un archivo
        nombre_archivo = 'resultados.docx'
        document.save(nombre_archivo)

        # Descargamos el archivo Word a la carpeta de descargas del usuario
        from google.colab import files
        files.download(nombre_archivo, root_dir='/content/downloads')

        st.success(f'Los resultados han sido exportados al archivo {nombre_archivo}')

